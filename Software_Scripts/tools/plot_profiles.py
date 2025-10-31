"""
This script gathers all the assigned profiles and generates
a plot file for each plot.  It also generates a csv for all
weekly and annual group profiles

All daily profiles referenced by the assigned profiles are
also plotted

Finally, a Word document is created that contains all the images
with corresponding profile name/ID

The images and word document are stored in the project "content" folder
"""
import iesve
import sys
import re
import os
import csv
import datetime as dt
from dateutil.relativedelta import relativedelta
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.ticker import MultipleLocator, FormatStrFormatter
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
from matplotlib.figure import Figure
import seaborn as sns


def print_profile(prof):
    s = "["
    for steps in list(prof):
        s += '[{:0.2f},{:0.2f},{}]'.format(steps.x, steps.y, steps.formula)
    s += "]"
    print(s)


def get_daily_profile_reference(id, day_profiles, proj):
    try:
        day_profile = day_profiles[id]
    except KeyError:
        day_profile = proj.daily_profile(id)
        if day_profile is not None:
            day_profiles[id] = day_profile

    if day_profile is None:
        return "Not Found"

    return day_profile.reference


def save_group_profile(prof, folder, project, day_profiles, category):
    profile_file_name = generate_filename(prof, folder) + ".csv"
    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday", "Holiday", "Heating-Rm", "Cooling-Rm", "Heating-Sys", "Cooling-Sys"]
    maxNumDays = len(days)

    if profile_file_name:
        profile_data = prof.get_data()
        if len(profile_data) >= 1:
            with open(profile_file_name, 'w', newline='') as csv_file:
                csv_writer = csv.writer(csv_file, dialect='excel', delimiter=',', quotechar='|', quoting=csv.QUOTE_MINIMAL)
                if prof.is_weekly():
                    # weekly profile
                    csv_writer.writerow(['Day', 'DailyID', 'Reference'])
                    for i, entry in enumerate(profile_data):
                        if i > maxNumDays:
                            break
                        csv_writer.writerow([days[i], entry, get_daily_profile_reference(entry, day_profiles, project)])
                elif prof.is_yearly():
                    # annual profile
                    csv_writer.writerow(['WeeklyID', 'Reference', 'DayFrom', 'DayTo', 'DateFrom', 'DateTo'])
                    for entry in profile_data:
                        fromDate = dt.date(1999, 1, 1) + relativedelta(days=+(entry[1] - 1))
                        to_date = dt.date(1999, 1, 1) + relativedelta(days=+(entry[2] - 1))
                        csv_writer.writerow([entry[0], 'TODO', entry[1], entry[2], fromDate.strftime('%d/%m'), to_date.strftime('%d/%m')])
            profile_name = '{0} [{1}]'.format(prof.reference, prof.id)
            project.register_content(profile_file_name, "{} Profiles (csv)".format(category), profile_name, False)


def ensure_dir(f):
    d = os.path.dirname(f)
    if not os.path.exists(d):
        os.makedirs(d)


def clean_dir(target_folder, proj):
    import os
    for the_file in os.listdir(target_folder):
        file_path = os.path.join(target_folder, the_file)
        try:
            if os.path.isfile(file_path):
                proj.deregister_content(file_path)
                os.remove(file_path)
            elif os.path.isdir(file_path):
                clean_dir(file_path, proj)
                os.rmdir(file_path)
        except:
            print("Failure to remove file: {0}".format(file_path))


def generate_filename(prof, folder):
    profile_name = prof.reference
    profile_name = re.sub(r'[<>:"/\|?*]', "", profile_name)
    type = "Group" if prof.is_group() else "Daily"
    profile_file_name = '{profile_name} [{type}-{id}]'.format(type=type, profile_name=profile_name, id=prof.id)
    full_path = folder + profile_file_name
    alt_path = folder + "Profile " + prof.id

    max_chars = 251  # Windows file limit minus file extension length (255 - 4 = 251)
    max_path = 255  # Total path limit minus file extension length (259 - 4 = 255)

    if len(profile_file_name) <= max_chars:
        if len(full_path) <= max_path:
            return full_path
        elif len(alt_path) <= max_path:
            print ("Warning: Profile name not included in filename because the total length would be too long. The total path length must be 259 characters or less.")
            return alt_path
        else:
            print ("Error: File not generated because the path length is too long. The total path length must be 259 characters or less.", file=sys.stderr)
            return ""
    else:
        if len(alt_path) <= max_path:
            print ("Warning: Profile name not included in filename because the total length would be too long. The total filename length must be 255 characters or less.")
            return alt_path
        else:
            print ("Error: File not generated because the path length is too long. The total path length must be 259 characters or less.", file=sys.stderr)
            return ""


def plot_day_profile_to_subplot(curax, prof):
    x_vals = []
    y_vals = []

    if prof is not None:
        profile_data = prof.get_data()
        for steps in profile_data:
            x_vals.append(steps[0])
            y_vals.append(steps[1])
        if len(profile_data) == 1:
            x_vals.append(24.0)
            y_vals.append(profile_data[0][1])
    else:
        print("Warning: no daily profile data to plot")

    curax.plot(x_vals, y_vals, color='#ffa500', lw=2)


def chart_week_group_profile(prof, target_folder, ve_proj, day_profiles, group_filenames):
    filename = generate_filename(prof, target_folder)

    if filename == "":
        return
    profile_name = '{0} [{1}]'.format(prof.reference, prof.id)

    print("Plotting weekly profile %s [%s]" % (prof.reference, prof.id))
    profile_data = prof.get_data()
    if len(profile_data) != 10:
        print ("Warning: skipping incomplete week profile:" + prof.id)
        return

    # create a suitable subplot layout
    cur_fig, ax = plt.subplots(1, 10, sharex=True, sharey=True, subplot_kw=dict(xlim=(0, 24)))
    cur_fig.set_size_inches(22.5, 3.0)

    cur_fig.suptitle('Weekly profile: ' + prof.reference + ' (' + prof.id + ')', fontsize=12, fontweight='bold')

    day_names = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday", "Holiday", "Heating", "Cooling"]
    num_days = len(day_names)
    major_locator = MultipleLocator(8)
    major_formatter = FormatStrFormatter('%d:00')
    minor_locator = MultipleLocator(4)

    for i, profile_entry in enumerate(profile_data):
        if i >= num_days:
            break

        try:
            day_profile = day_profiles[profile_entry]
        except KeyError:
            day_profile = ve_proj.daily_profile(profile_entry)
            if day_profile is not None:
                day_profiles[profile_entry] = day_profile

        if day_profile is None:
            print("Warning: daily profile not found for ID: %s (weekly profile ID: %s)" % (profile_entry, prof.id))

        current_subplot = ax[i]

        plot_day_profile_to_subplot(current_subplot, day_profile)
        current_subplot.set_title(day_names[i])

        current_subplot.xaxis.set_major_locator(major_locator)
        current_subplot.xaxis.set_major_formatter(major_formatter)

        # For the minor ticks, use no labels; default NullFormatter
        current_subplot.xaxis.set_minor_locator(minor_locator)

        if i == 0:
            # set a label for the first plot only
            if prof.is_absolute():
                current_subplot.set_ylabel("Absolute value")
            else:
                current_subplot.set_ylabel("Modulating value")

    sns.despine()
    plt.tight_layout()   # tight_layout helps with the overwriting of string but kinda messes up the axis scales

    cur_fig.subplots_adjust(wspace=0.3, left=0.04, right=0.98, top=0.85)

    filename_png = filename + '.png'
    cur_fig.savefig(filename_png)
    group_filenames[prof.id] = filename_png

    ve_proj.register_content(filename_png, "Weekly Profiles", profile_name, False)

    cur_fig.clf()
    plt.clf()
    plt.close('all')


def chart_annual_group_profile(prof, target_folder, ve_proj, group_profiles, day_profiles, group_filenames):
    filename = generate_filename(prof, target_folder)

    if filename == "":
        return

    profile_name = '{0} [{1}]'.format(prof.reference, prof.id)

    print("Plotting annual profile %s [%s]" % (prof.reference, prof.id))
    profile_data = prof.get_data()
    num_entries = len(profile_data)
    day_names = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday", "Holiday", "Heating-Rm", "Cooling-Rm", "Heating-Sys", "Cooling-Sys"]
    num_days = len(day_names)

    # create a suitable subplot layout
    cur_fig, ax = plt.subplots(num_entries, num_days, sharex=True, sharey=True, subplot_kw=dict(xlim=(0, 24)))
    cur_fig.set_size_inches(22, (num_entries * 2) + 1)  # a plot is 2-ish" high, and 7-ish across for 8

    major_locator = MultipleLocator(8)
    major_formatter = FormatStrFormatter('%d:00')
    minor_locator = MultipleLocator(4)

    for weeks, annual_profile_entry in enumerate(profile_data):
        try:
            week_profile = group_profiles[annual_profile_entry[0]]
        except KeyError:
            week_profile = ve_proj.groupProfile(annual_profile_entry[0])
            if week_profile is not None:
                group_profiles[annual_profile_entry.id] = week_profile

        if week_profile is None:
            print("Warning: weekly profile not found for ID: %s (annual profile ID: %s)" % (annual_profile_entry[0], prof.id))

        if week_profile is not None:
            week_profile_data = week_profile.get_data()

            for days, day_profile_entry in enumerate(week_profile_data):
                if num_entries == 1:
                    current_subplot = ax[days]
                else:
                    current_subplot = ax[weeks, days]

                try:
                    day_profile = day_profiles[day_profile_entry]
                except KeyError:
                    day_profile = ve_proj.daily_profile(day_profile_entry)
                    if day_profile is not None:
                        day_profiles[day_profile_entry] = day_profile

                if day_profile is None:
                    print("Warning: daily profile %d not found for ID: %s (annual profile ID: %s, weekly Profile ID: %s)"
                          % (days, day_profile_entry, prof.id, annual_profile_entry[0]))

                plot_day_profile_to_subplot(current_subplot, day_profile)

                current_subplot.xaxis.set_major_locator(major_locator)
                current_subplot.xaxis.set_major_formatter(major_formatter)

                # For the minor ticks, use no labels; default NullFormatter
                current_subplot.xaxis.set_minor_locator(minor_locator)

                # for every day, write a title (effectively a column header) but only for the first row
                if weeks == 0:
                    current_subplot.set_title(day_names[days])

                # for every row, write the dates covered by this profile (only write for first column)
                if days == 0:
                    # set a label for the first plot only
                    fromDate = dt.date(1999, 1, 1)+relativedelta(days=+(annual_profile_entry[1] - 1))
                    to_date = dt.date(1999, 1, 1)+relativedelta(days=+(annual_profile_entry[2] - 1))
                    current_subplot.set_ylabel(fromDate.strftime('%d/%m') + ' - ' + to_date.strftime('%d/%m'))

    # now tighten up the spacing and save out the chart to disk
    plt.tight_layout()   # tight_layout helps with the overwriting of string but kinda messes up the axis scales

    # leave 20 pixels space for the title - this needs to be calculated as a percentage for subplots_adjust
    # assume 300dpi, so 60 pixels for the title
    total_height_DPI = 100.0 * ((num_entries * 2.0) + 1.0)
    ratio = 1.0 - ((40.0 + (6.0 * num_entries)) / total_height_DPI)
    cur_fig.subplots_adjust(top=ratio)

    cur_fig.suptitle('Annual profile: ' + prof.reference + ' (' + prof.id + ')', fontsize=12, fontweight='bold')
    sns.despine()

    filename_png = filename + '.png'
    cur_fig.savefig(filename_png)
    group_filenames[prof.id] = filename_png

    ve_proj.register_content(filename_png, "Annual profiles", profile_name, False)

    cur_fig.clf()
    plt.clf()
    plt.close('all')


def chart_compact_profile(prof, target_folder, ve_proj, group_filenames):
    filename = generate_filename(prof, target_folder)

    if filename == "":
        return

    profile_name = '{0} [{1}]'.format(prof.reference, prof.id)

    print("Plotting compact profile %s [%s]" % (prof.reference, prof.id))
    data = prof.get_data()
    num_entries = len(data)

    # create a suitable subplot layout
    cur_fig, ax = plt.subplots(num_entries, 1, sharex=False, sharey=True, subplot_kw=dict(xlim=(0, 24)))
    cur_fig.set_size_inches(8, (num_entries * 3) + 1)  # a plot is 3-ish" high, and 8-ish across for the bar plot

    major_locator = MultipleLocator(4)
    major_formatter = FormatStrFormatter('%d:00')
    minor_locator = MultipleLocator(1)

    print("Compact profile: num_entries = %d" % (num_entries))

    for entry in range(0, num_entries):
        profile_entry = data[entry]
        end_day, end_month = profile_entry[0]
        to_date = dt.date(1999, end_month, end_day)

        # subplots returns an array if num_entries > 1, so check if
        # we need to fetch the sub-plot out of the array
        current_subplot = ax if num_entries == 1 else ax[entry]
        y_vals = []
        y_labels = []
        y_val = 10
        for period in reversed(profile_entry[1:]):
            x_vals = []
            for tp in period[1:]:
                if tp[0] is False:
                    continue        # don't bother with switched off time periods
                hour_on, minute_on, hourOff, minute_off = tp[1:]
                time_on = hour_on + (minute_on / 60.0)
                time_off = hourOff + (minute_off / 60.0)

                # a compact profile is always on if the two times are the same
                if time_on == time_off:
                    time_on = 0.0
                    time_off = 24.0

                if time_off == 0:
                    time_off = 24.0
                if time_off < time_on:
                    print("Warning: invalid compact profile entry, time_off before time_on")
                    time_off = time_on

                x_vals.append((time_on, time_off - time_on))

            height_vals = (y_val - 3, 6)
            current_subplot.broken_barh(x_vals, height_vals, facecolors='#ffa500')
            y_vals.append(y_val)
            y_labels.append(period[0])
            y_val = y_val + 10

        current_subplot.yaxis.set_ticks(y_vals)
        current_subplot.yaxis.set_ticklabels(y_labels)
        current_subplot.set_ylim(0, (10 * (len(y_labels)) + 5))

        current_subplot.xaxis.set_major_locator(major_locator)
        current_subplot.xaxis.set_major_formatter(major_formatter)
        current_subplot.xaxis.grid(True, linestyle='dotted', alpha=0.5, color='b')

        # For the minor ticks, use no labels; default NullFormatter
        current_subplot.xaxis.set_minor_locator(minor_locator)

        # for every entry, write a title
        current_subplot.text(1.0, 1.0, 'End date: {date}'.format(date=to_date.strftime('%d/%m')), ha='left', va='bottom', size='medium')

    # now tighten up the spacing and save out the chart to disk
    plt.tight_layout()   # tight_layout helps with the overwriting of string but kinda messes up the axis scales

    # leave 20 pixels space for the title - this needs to be calculated as a percentage for subplots_adjust
    # assume 300dpi, so 60 pixels for the title
    total_height_DPI = 100.0 * ((num_entries * 2.0) + 1.0)
    ratio = 1.0 - ((40.0 + (6.0 * num_entries)) / total_height_DPI)
    cur_fig.subplots_adjust(top=ratio)

    cur_fig.suptitle('Compact profile: ' + prof.reference + ' (' + prof.id + ')', fontsize=12, fontweight='bold')
    sns.despine()

    filename_png = filename + '.png'
    cur_fig.savefig(filename_png)
    group_filenames[prof.id] = filename_png

    ve_proj.register_content(filename_png, "Compact profiles", profile_name, False)

    cur_fig.clf()
    plt.clf()
    plt.close('all')


def chart_free_form_profile(prof, target_folder, ve_proj, group_filenames):
    formatter = mdates.DateFormatter('%b')
    monday = 1
    mondays = mdates.WeekdayLocator(monday)
    locator = mdates.MonthLocator()

    filename = generate_filename(prof, target_folder)

    if filename == "":
        return

    print("Plotting free-form profile %s [%s]" % (prof.reference, prof.id))
    fig = Figure()
    FigureCanvas(fig)
    fig.add_subplot(111)  # "111" means "1x1 grid, first subplot"

    x_vals = []
    y_vals = []

    if prof is not None:
        profile_data = prof.get_data()
        for entry in list(profile_data):
            x_vals.append(mdates.date2num(dt.datetime(2015, entry[0], entry[1], entry[2], entry[3])))
            y_vals.append(entry[4])
    else:
        print("Warning: no daily profile data to plot")

    plt.plot(x_vals, y_vals, color='#ffa500', linewidth=0.5)

    cur_axis = plt.gca()
    cur_axis.xaxis.set_major_locator(locator)
    cur_axis.xaxis.set_major_formatter(formatter)
    cur_axis.xaxis.set_minor_locator(mondays)
    cur_axis.autoscale_view()

    plt.suptitle(prof.reference + ' (' + prof.id + ')')

    for label in cur_axis.get_xmajorticklabels():
        label.set_fontsize(7)

    # Set the axis labels
    plt.xlabel("Time")
    if prof.is_absolute():
        plt.ylabel("Absolute value")
    else:
        plt.ylabel("Modulating value")

    sns.despine()

    fig.autofmt_xdate()

    filename_png = filename + '.png'
    plt.savefig(filename_png)
    group_filenames[prof.id] = filename_png

    profile_name = '{0} [{1}]'.format(prof.reference, prof.id)
    ve_proj.register_content(filename_png, "Free form profiles", profile_name, False)

    plt.clf()
    fig.clf()
    plt.close('all')


def create_chart(prof, target_folder, chart_dict, project):
    filename = generate_filename(prof, target_folder)

    if filename == "":
        return

    profile_name = '{0} [{1}]'.format(prof.reference, prof.id)

    print("Plotting daily profile %s [%s]" % (prof.reference, prof.id))
    fig = Figure()
    FigureCanvas(fig)
    fig.add_subplot(111)  # "111" means "1x1 grid, first subplot"

    x_vals = []
    y_vals = []

    if prof is not None:
        profile_data = prof.get_data()
        for steps in profile_data:
            x_vals.append(steps[0])
            y_vals.append(steps[1])
        if len(profile_data) == 1:
            x_vals.append(24.0)
            y_vals.append(profile_data[0][1])
    else:
        print("Warning: no daily profile data to plot")

    major_locator = MultipleLocator(4)
    major_formatter = FormatStrFormatter('%d:00')
    minor_locator = MultipleLocator(1)

    plt.plot(x_vals, y_vals, color='#ffa500', lw=2)

    cur_axis = plt.gca()
    cur_axis.xaxis.set_major_locator(major_locator)
    cur_axis.xaxis.set_major_formatter(major_formatter)

    # For the minor ticks, use no labels; default NullFormatter
    cur_axis.xaxis.set_minor_locator(minor_locator)

    plt.grid(True, 'major')
    plt.suptitle(prof.reference + ' (' + prof.id + ')')
    plt.xlim(0, 24)

    # Set the axis labels
    plt.xlabel("Time of Day")
    if prof.is_absolute():
        plt.ylabel("Absolute value")
    else:
        plt.ylabel("Modulating value")

    sns.despine()

    filename_png = filename + '.png'
    plt.savefig(filename_png)
    chart_dict[prof.id] = filename_png

    project.register_content(filename_png, "Daily profiles", profile_name, False)

    plt.clf()
    fig.clf()
    plt.close('all')


def create_word_index_document(building_target_folder, building_reference, proj, day_profiles, group_profiles, group_filenames, chart_dict):
    from docx import Document
    from docx.shared import Inches
    import collections

    filename = proj.name
    if building_reference is not None:
        filename += ' ({})'.format(building_reference)
    filename += ' profiles.docx'
    filename = re.sub(r'[<>:"/\|?*]', "", filename)

    document = Document()
    filename = building_target_folder + filename

    print("Writing word index document")

    document.core_properties.author = "IES Virtual Environment"
    document.core_properties.created = dt.datetime.now()

    image_width = Inches(3.5)

    if len(group_filenames) > 0:
        document.add_heading('Model assigned profiles', 0)
        group_table = document.add_table(rows=1, cols=3)
        group_table.style = 'Light Grid Accent 1'
        hdr_cells = group_table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Desc'
        hdr_cells[2].text = 'Profile'
        od = collections.OrderedDict(sorted(group_filenames.items()))
        for key, value in od.items():
            row_cells = group_table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = group_profiles[key].reference

            cell = row_cells[2]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(value, width=image_width)

    if len(chart_dict) > 0:
        # if we also plotted group profiles, then introduce a page break for clarity
        if len(group_filenames) > 0:
            document.add_page_break()

        document.add_heading('Daily profiles detailed plot', 0)
        day_table = document.add_table(rows=1, cols=3)
        day_table.style = 'Light Grid Accent 3'
        hdr_cells = day_table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Desc'
        hdr_cells[2].text = 'Profile'
        od = collections.OrderedDict(sorted(chart_dict.items()))
        for key, value in od.items():
            row_cells = day_table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = day_profiles[key].reference

            cell = row_cells[2]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(value, width=image_width)

    document.save(filename)
    proj.register_content(filename, 'Assigned profiles', 'Assigned profiles', True)

# # MAIN BODY STARTS HERE # #
if __name__ == '__main__':
    # clear any old plotting data, also in case a previous plot was cancelled / interrupted
    plt.clf()
    plt.close('all')
    sns.set_style("ticks")

    # now get the project data
    proj = iesve.VEProject.get_current_project()
    target_folder = proj.content_folder + '\\Profiles\\'
    print("Writing profiles to folder: " + target_folder)
    ensure_dir(target_folder)
    clean_dir(target_folder, proj)
    model_list = proj.models
    print("Number of building models: %d" % (len(model_list)))
    for building in model_list:
        print("Building: %s" % (building.id))
        building_target_folder = target_folder + building.id + '\\'
        ensure_dir(building_target_folder)

        group_filenames = {}
        chart_dict = {}

        # get all the assigned profiles in the building, so that we can plot them
        assigned_profile_ids = building.get_assigned_profiles()
        group_profiles = proj.group_profiles(assigned_profile_ids)
        day_profiles = {}

        # alternatively - to plot all defined profiles, use the following (but don't iterate over all buildings)
        # day_profiles, group_profiles = proj.profiles()

        # create a list of IDs to process - we store these as a list so that we can
        # append new items as we find dependencies
        list_of_IDs = []
        for profile_id, a_profile in group_profiles.items():
            list_of_IDs.append(profile_id)

        print("Number of group profiles: %d" % (len(list_of_IDs)))

        # now chart of the group profiles - because these also reference the day profiles
        # we pass in all the daily profiles we have already fetched so they can be re-used
        for profile_id in list_of_IDs:
            try:
                a_profile = group_profiles[profile_id]
            except:
                print("Error: unable to find profile with ID: {}".format(profile_id))
                continue

            if a_profile.is_weekly():
                save_group_profile(a_profile, building_target_folder, proj, day_profiles, "Weekly")
                chart_week_group_profile(a_profile, building_target_folder, proj, day_profiles, group_filenames)
            elif a_profile.is_yearly():
                save_group_profile(a_profile, building_target_folder, proj, day_profiles, "Annual")
                chart_annual_group_profile(a_profile, building_target_folder, proj, group_profiles, day_profiles, group_filenames)
            elif a_profile.is_compact():
                chart_compact_profile(a_profile, building_target_folder, proj, group_filenames)
            elif a_profile.is_freeform():
                assert isinstance(a_profile, iesve.FreeFormProfile), "Profile is not Free-form"
                a_profile.load_data()
                if a_profile.is_graphable():
                    chart_free_form_profile(a_profile, building_target_folder, proj, group_filenames)

        # optional - chart all the daily profiles again in all detail
        print("Number of derived daily profiles: %d" % (len(day_profiles)))
        for day_profile_id, day_profile in day_profiles.items():
            create_chart(day_profile, building_target_folder, chart_dict, proj)

        building_reference = None
        if len(model_list) > 1:
            building_reference = building.id

        create_word_index_document(building_target_folder, building_reference, proj, day_profiles, group_profiles, group_filenames, chart_dict)

        del day_profiles
        del group_profiles

    print("*** done ***")
